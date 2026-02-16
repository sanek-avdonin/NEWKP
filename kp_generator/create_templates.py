"""
Создаёт примеры шаблонов КП в assets/templates/.
Запуск из корня проекта: python -m kp_generator.create_templates
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook

# Путь к папке шаблонов (рядом с create_templates.py -> assets/templates)
BASE = Path(__file__).resolve().parent
TEMPLATES_DIR = BASE / "assets" / "templates"


def create_docx_template() -> Path:
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    out = TEMPLATES_DIR / "template_kp.docx"

    doc = Document()
    doc.add_paragraph("Коммерческое предложение").alignment = 1  # center
    doc.add_paragraph()
    doc.add_paragraph("Компания: {{COMPANY_NAME}}")
    doc.add_paragraph("ИНН: {{INN}}")
    doc.add_paragraph("Адрес: {{ADDRESS}}")
    doc.add_paragraph("Телефон: {{PHONE}}")
    doc.add_paragraph("Руководитель: {{CEO}}")
    doc.add_paragraph()

    # Таблица: заголовки + строка-образец + итого
    table = doc.add_table(rows=4, cols=5)
    table.style = "Table Grid"
    headers = ["№", "Наименование", "Количество", "Ед. изм.", "Цена, руб", "Сумма, руб"]
    # 5 колонок без № или 6 с № - в HEADER_SYNONYMS нет "№", делаем 5: Наименование, Количество, Ед. изм., Цена, Сумма
    h = ["Наименование", "Количество", "Ед. изм.", "Цена, руб", "Сумма, руб"]
    for i, text in enumerate(h):
        table.rows[0].cells[i].text = text
    # Строка-образец
    table.rows[1].cells[0].text = "Пример товара"
    table.rows[1].cells[1].text = "1"
    table.rows[1].cells[2].text = "шт"
    table.rows[1].cells[3].text = "100,00"
    table.rows[1].cells[4].text = "100,00"
    # Пустая строка (будет удалена/заменена при рендере)
    table.rows[2].cells[0].text = ""
    table.rows[2].cells[1].text = ""
    table.rows[2].cells[2].text = ""
    table.rows[2].cells[3].text = ""
    table.rows[2].cells[4].text = ""
    # Итого
    table.rows[3].cells[0].text = "Итого"
    table.rows[3].cells[1].text = ""
    table.rows[3].cells[2].text = ""
    table.rows[3].cells[3].text = ""
    table.rows[3].cells[4].text = "0,00"

    doc.save(out)
    return out


def create_excel_template() -> Path:
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    out = TEMPLATES_DIR / "template_kp.xlsx"

    wb = Workbook()
    ws = wb.active
    ws.title = "КП"

    ws["A1"] = "Коммерческое предложение"
    ws["A3"] = "Компания: {{COMPANY_NAME}}"
    ws["A4"] = "ИНН: {{INN}}"
    ws["A5"] = "Адрес: {{ADDRESS}}"
    ws["A6"] = "Телефон: {{PHONE}}"
    ws["A7"] = "Руководитель: {{CEO}}"
    ws["A9"] = ""

    header_row = 10
    for c, h in enumerate(["Наименование", "Количество", "Ед. измерения", "Цена", "Сумма"], start=1):
        ws.cell(header_row, c).value = h
    ws.cell(11, 1).value = "Пример позиции"
    ws.cell(11, 2).value = 1
    ws.cell(11, 3).value = "шт"
    ws.cell(11, 4).value = 100.0
    ws.cell(11, 5).value = 100.0
    ws.cell(12, 1).value = "Итого"
    ws.cell(12, 5).value = 0.0

    wb.save(out)
    return out


def main():
    print("Создание примеров шаблонов в", TEMPLATES_DIR)
    p1 = create_docx_template()
    p2 = create_excel_template()
    print("Создано:", p1, p2)


if __name__ == "__main__":
    main()
