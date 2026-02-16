from __future__ import annotations

from copy import deepcopy
from decimal import Decimal, ROUND_HALF_UP
from typing import Dict, List, Optional, Tuple
import re

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from kp_generator.models import CompanyProfile, Item  # абсолютный импорт


# ---------------------------
# 1) Placeholders for company
# ---------------------------

PLACEHOLDERS: Dict[str, str] = {
    "{{COMPANY_NAME}}": "name",
    "{{INN}}": "inn",
    "{{ADDRESS}}": "address",
    "{{PHONE}}": "phone",
    "{{CEO}}": "ceo",
}


# ---------------------------
# 2) Header synonyms
# ---------------------------

HEADER_SYNONYMS = {
    "name": ["наименование", "товар", "услуга", "описание объекта закупки", "наименование товара"],
    "qty": ["количество", "кол-во", "кол во", "кол.", "qty"],
    "unit": ["ед", "ед.", "едизм", "ед.изм", "единица измерения"],
    "price": ["цена", "цена за ед", "цена за единицу", "цена, руб"],
    "amount": ["сумма", "стоимость", "итого", "сумма, руб"],
}


# ---------------------------
# 3) Text / number utilities
# ---------------------------

def _normalize_text(s: str) -> str:
    """Lowercase + collapse spaces + strip."""
    s = (s or "").replace("\u00a0", " ")
    s = re.sub(r"\s+", " ", s).strip().lower()
    return s


def _fmt_money(v: Decimal) -> str:
    """Format Decimal to '12 345,67'."""
    q = v.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    # Decimal to string with dot
    s = f"{q:.2f}".replace(".", ",")
    int_part, frac_part = s.split(",")

    # thousand separators with space
    sign = ""
    if int_part.startswith("-"):
        sign = "-"
        int_part = int_part[1:]

    int_part = "{:,}".format(int(int_part)).replace(",", " ")
    return f"{sign}{int_part},{frac_part}"


def _to_decimal_safe(s: str) -> Decimal:
    """Parse numeric string with RU format to Decimal (best-effort)."""
    s = (s or "").replace("\u00a0", " ").strip()
    s = s.replace(" ", "").replace(",", ".")
    s = re.sub(r"[^\d.\-]", "", s)
    if s == "" or s == "-" or s == ".":
        return Decimal("0")
    try:
        return Decimal(s)
    except Exception:
        return Decimal("0")


def _fmt_qty(q: Decimal) -> str:
    """Qty: integer without decimals, else with comma."""
    if q == q.to_integral():
        return str(int(q))
    # normalize to remove exponent etc.
    s = format(q.normalize(), "f")
    return s.replace(".", ",")


# ---------------------------
# 4) Safe text set in paragraph/cell
# ---------------------------

def _set_paragraph_text_preserve_runs(p: Paragraph, new_text: str) -> None:
    """
    Replace paragraph text while preserving formatting as much as possible:
    - take all runs -> one string replaced -> write to first run, clear the rest.
    """
    if not p.runs:
        p.add_run(new_text)
        return

    # clear all runs, set in first run
    for r in p.runs:
        r.text = ""
    p.runs[0].text = new_text


def _set_cell_text_preserve(cell: _Cell, new_text: str) -> None:
    """
    Set cell text without resetting table formatting:
    update first paragraph runs instead of cell.text assignment.
    """
    if cell.paragraphs:
        _set_paragraph_text_preserve_runs(cell.paragraphs[0], new_text)
        # clear extra paragraphs (optional, conservative):
        for extra_p in cell.paragraphs[1:]:
            _set_paragraph_text_preserve_runs(extra_p, "")
    else:
        p = cell.add_paragraph()
        _set_paragraph_text_preserve_runs(p, new_text)


# ---------------------------
# 5) Placeholder replacement
# ---------------------------

def _replace_placeholders_in_paragraph(p: Paragraph, company: CompanyProfile) -> None:
    if not p.runs:
        return

    full = "".join(r.text for r in p.runs)
    if "{{" not in full or "}}" not in full:
        return

    for ph, attr in PLACEHOLDERS.items():
        if ph in full:
            full = full.replace(ph, str(getattr(company, attr, "")))

    _set_paragraph_text_preserve_runs(p, full)


def _replace_placeholders_in_doc(doc: Document, company: CompanyProfile) -> None:
    # paragraphs
    for p in doc.paragraphs:
        _replace_placeholders_in_paragraph(p, company)

    # tables
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_placeholders_in_paragraph(p, company)


# ---------------------------
# 6) Table/header detection
# ---------------------------

def _find_header_row_and_colmap(table: Table, max_scan_rows: int = 5) -> Tuple[int, Dict[str, int], int]:
    """
    Find goods table header and column map.
    We scan first N rows and build col_map possibly from multiple header rows.

    Returns:
      header_end_row_idx (int): last header row index (so sample row is header_end+1)
      col_map (dict): field -> column index
      header_start_row_idx (int): first row index where header clues were found (for debugging)
    """
    scan_rows = min(max_scan_rows, len(table.rows))
    col_map: Dict[str, int] = {}
    header_rows_found: List[int] = []

    for r in range(scan_rows):
        row = table.rows[r]
        for c_idx, cell in enumerate(row.cells):
            txt = _normalize_text(cell.text)
            if not txt:
                continue

            for key, syns in HEADER_SYNONYMS.items():
                if key in col_map:
                    continue
                if any(syn in txt for syn in syns):
                    col_map[key] = c_idx
                    header_rows_found.append(r)

    # minimal required: name, qty, and at least one of price/amount
    if "name" in col_map and "qty" in col_map and ("price" in col_map or "amount" in col_map):
        header_start = min(header_rows_found) if header_rows_found else 0
        header_end = max(header_rows_found) if header_rows_found else 0
        return header_end, col_map, header_start

    raise ValueError(
        "Не удалось найти таблицу товаров: не обнаружены заголовки колонок "
        "(требуется минимум: 'name', 'qty' и 'price' или 'amount')."
    )


def _find_goods_table(doc: Document) -> Tuple[Table, int, Dict[str, int]]:
    """
    Iterate over all tables and return first matching goods table.
    Returns (table, header_end_row_idx, col_map).
    """
    for t in doc.tables:
        try:
            header_end, col_map, _ = _find_header_row_and_colmap(t, max_scan_rows=5)
            return t, header_end, col_map
        except ValueError:
            continue

    raise ValueError(
        "Не удалось найти таблицу с товарами в документе. "
        "Проверьте, что в одной из таблиц есть заголовки колонок (Наименование/Количество/Цена/Сумма и т.п.)."
    )


# ---------------------------
# 7) Totals detection
# ---------------------------

def _classify_total_row(left_text: str) -> Optional[str]:
    """
    Determine row type by keywords.
    Returns: 'total', 'vat', 'total_without_vat' or None.
    """
    t = _normalize_text(left_text)
    if not t:
        return None

    # priority: VAT markers first
    if "без ндс" in t or "итого без ндс" in t:
        return "total_without_vat"
    if "ндс" in t or "в том числе ндс" in t:
        return "vat"
    if "итого" in t or "всего" in t:
        return "total"
    return None


def _find_total_rows(table: Table, start_row: int) -> Dict[str, int]:
    """
    Find total rows in table starting from row index start_row.
    Returns mapping: type -> row_idx (first occurrence).
    """
    found: Dict[str, int] = {}
    for i in range(start_row, len(table.rows)):
        row = table.rows[i]
        if not row.cells:
            continue
        left = row.cells[0].text
        row_type = _classify_total_row(left)
        if row_type and row_type not in found:
            found[row_type] = i
    return found


# ---------------------------
# 8) Row deletion / cloning
# ---------------------------

def _delete_rows(table: Table, indices: List[int]) -> None:
    """Delete table rows by indices (reverse order)."""
    for i in sorted(set(indices), reverse=True):
        if 0 <= i < len(table.rows):
            table._tbl.remove(table.rows[i]._tr)


def _clone_and_insert_row_before_tr(table: Table, src_row_idx: int, target_tr) -> None:
    """
    Clone src row as XML and insert it before target_tr.
    """
    src_tr = table.rows[src_row_idx]._tr
    new_tr = deepcopy(src_tr)
    target_tr.addprevious(new_tr)


def _clone_and_append_row(table: Table, src_row_idx: int) -> None:
    """Clone src row as XML and append to end of table."""
    src_tr = table.rows[src_row_idx]._tr
    new_tr = deepcopy(src_tr)
    table._tbl.append(new_tr)


# ---------------------------
# 9) Fill row values
# ---------------------------

def _fill_row_cells(row_cells: List[_Cell], col_map: Dict[str, int], item_no: int, it: Item) -> None:
    """
    Fill row cells by mapping. We don't require a placeholder column for row number;
    if template has a first column '№' and it's not mapped, we still can fill if it's before name.
    Here: if we see a 'name' col and there's a col before it, we optionally fill it as row number
    ONLY if that first col looks like number column (empty or digits).
    """
    # optionally fill "№" column if exists (heuristic)
    name_col = col_map.get("name")
    if name_col is not None and name_col > 0:
        first_cell_txt = _normalize_text(row_cells[0].text)
        # if empty or numeric-like, consider it as row number col
        if first_cell_txt == "" or re.fullmatch(r"\d+", first_cell_txt or ""):
            _set_cell_text_preserve(row_cells[0], str(item_no))

    if "name" in col_map:
        _set_cell_text_preserve(row_cells[col_map["name"]], it.name)

    if "qty" in col_map:
        _set_cell_text_preserve(row_cells[col_map["qty"]], _fmt_qty(it.qty))

    if "unit" in col_map:
        _set_cell_text_preserve(row_cells[col_map["unit"]], it.unit or "")

    if "price" in col_map:
        _set_cell_text_preserve(row_cells[col_map["price"]], _fmt_money(it.price))

    if "amount" in col_map:
        _set_cell_text_preserve(row_cells[col_map["amount"]], _fmt_money(it.amount))


# ---------------------------
# 10) Totals update
# ---------------------------

def _update_totals(
    table: Table,
    col_map: Dict[str, int],
    total_rows_idx: Dict[str, int],
    index_shift: int,
    items: List[Item]
) -> None:
    """
    Update totals rows after deletion/insertion.
    index_shift = (inserted_count - deleted_count_between_sample_and_total)
    So actual row index = saved_index + index_shift
    """
    # choose amount column
    amount_col = col_map.get("amount") or col_map.get("price")
    if amount_col is None:
        return

    total_sum = sum((it.amount for it in items), Decimal("0.00"))

    # VAT 20% included: vat = total*20/120
    vat = (total_sum * Decimal("20") / Decimal("120")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    total_wo_vat = (total_sum - vat).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)

    def safe_set(row_idx: int, value: Decimal) -> None:
        if 0 <= row_idx < len(table.rows):
            row = table.rows[row_idx]
            if amount_col < len(row.cells):
                _set_cell_text_preserve(row.cells[amount_col], _fmt_money(value))

    for typ, old_idx in total_rows_idx.items():
        new_idx = old_idx + index_shift
        if typ == "total":
            safe_set(new_idx, total_sum)
        elif typ == "vat":
            safe_set(new_idx, vat)
        elif typ == "total_without_vat":
            safe_set(new_idx, total_wo_vat)


# ---------------------------
# 11) Main render function
# ---------------------------

def render_kp_docx(template_path: str, company: CompanyProfile, items: List[Item], output_path: str) -> None:
    """
    Render KP docx by:
      1) load template
      2) replace company placeholders
      3) find goods table + header col map
      4) determine sample row
      5) find totals rows
      6) delete old rows between sample and totals
      7) clone sample row for each item and insert before totals (or append)
      8) delete sample row
      9) update totals
      10) save
    """
    doc = Document(template_path)

    # Step 2: placeholders
    _replace_placeholders_in_doc(doc, company)

    # Step 3: find table + columns
    table, header_end_row, col_map = _find_goods_table(doc)

    # Step 4: sample row
    sample_row_idx = header_end_row + 1
    if sample_row_idx >= len(table.rows):
        raise ValueError(
            "В шаблоне найдены заголовки таблицы, но отсутствует строка-образец (строка сразу после заголовков). "
            "Добавьте хотя бы одну строку с примером оформления."
        )

    # Step 5: total rows (search after header)
    total_rows = _find_total_rows(table, start_row=sample_row_idx)
    first_total_idx = min(total_rows.values()) if total_rows else None

    # Step 6: delete old goods rows BETWEEN sample and first total (or to end)
    deleted_indices: List[int] = []
    if first_total_idx is not None:
        # delete sample+1 ... first_total-1
        deleted_indices = list(range(sample_row_idx + 1, first_total_idx))
    else:
        # delete sample+1 ... end-1
        deleted_indices = list(range(sample_row_idx + 1, len(table.rows)))

    deleted_count = len(deleted_indices)
    if deleted_indices:
        _delete_rows(table, deleted_indices)

    # After deletion, total rows moved UP by deleted_count if they were after deletion range.
    # We'll compute shift later precisely as (inserted_count - deleted_count).
    # But note: if totals exist, their saved indices are based on PRE-deletion state.
    # We didn't recompute totals indices; we will apply index_shift = inserted_count - deleted_count
    # (since deletion moves them up, insertion moves them down).
    # This works because we delete only rows strictly before first total.
    # (If template has totals further down, also before them — still valid)

    # Determine insertion target TR (first total row) AFTER deletion
    inserted_count = 0
    if first_total_idx is not None:
        # first_total_idx in pre-deletion indices; after deletion it becomes:
        first_total_after_delete = first_total_idx - deleted_count
        # guard
        if first_total_after_delete < 0 or first_total_after_delete >= len(table.rows):
            # fallback: recompute totals after deletion
            total_rows_after = _find_total_rows(table, start_row=sample_row_idx)
            if total_rows_after:
                first_total_after_delete = min(total_rows_after.values())
            else:
                first_total_after_delete = None

        target_tr = table.rows[first_total_after_delete]._tr if first_total_after_delete is not None else None

        # Step 7: clone + insert before totals
        for idx, it in enumerate(items, start=1):
            if target_tr is not None:
                _clone_and_insert_row_before_tr(table, sample_row_idx, target_tr)
                # inserted row will appear immediately before target_tr
                inserted_row_idx = first_total_after_delete - 1  # new row inserted right before totals
            else:
                _clone_and_append_row(table, sample_row_idx)
                inserted_row_idx = len(table.rows) - 1

            inserted_count += 1

            # Fill inserted row
            row = table.rows[inserted_row_idx]
            _fill_row_cells(row.cells, col_map, idx, it)

        # Step 7 end
    else:
        # No totals -> append rows to end
        for idx, it in enumerate(items, start=1):
            _clone_and_append_row(table, sample_row_idx)
            inserted_count += 1
            row = table.rows[len(table.rows) - 1]
            _fill_row_cells(row.cells, col_map, idx, it)

    # Step 8: After inserting new rows, delete sample row (original)
    # Sample row index still valid because we only inserted after it (before totals or at end).
    _delete_rows(table, [sample_row_idx])

    # Step 9: update totals (indices: after deletion of data rows + insertion, then -1 for sample row removal)
    if total_rows:
        index_shift = inserted_count - deleted_count - 1  # -1 for sample row deletion
        _update_totals(table, col_map, total_rows, index_shift=index_shift, items=items)

    # Save
    doc.save(output_path)
